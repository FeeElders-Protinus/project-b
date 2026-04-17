import json
import re
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from openai import OpenAI
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── colours ──────────────────────────────────────────────────────────────────
COLOR_HEADER = "1F4E79"   # dark blue header row
COLOR_ROW_ALT = "EBF3FB"  # light-blue alternating rows

SYSTEM_PROMPT_NARRATIVE = """Je bent een ervaren IT-adviseur die professionele rapportages schrijft over de installed base van klanten.

Je ontvangt:
1. De doelen van de klant
2. Dashboard-statistieken (percentages en aantallen per producttype, regio en categorie)
3. Lijsten van closed-source en niet-EU producten
4. Per categorie een overzicht van vendors en hun producten

Schrijf een zakelijke analyse in het Nederlands met de volgende secties.
De alternatieven-tabellen (open source en EU) worden apart gegenereerd en in het rapport ingevoegd — schrijf die dus NIET zelf uit.

**Samenvatting**
Kort overzicht van de huidige situatie ten opzichte van alle gestelde doelen.

**Analyse per doel**
Behandel elk doel afzonderlijk:
- Huidige stand (concreet percentage/aantal uit de data)
- Of het doel behaald is of wat de kloof is
- Verwijzing dat de aanbevolen alternatieven direct onder deze sectie staan

**Vendor Consolidatie**
Per productcategorie met meerdere vendors: benoem de overlap en geef een concreet consolidatieadvies.

**Aanbevelingen**
Sluit af met 3-5 geprioriteerde, concrete actiepunten.

Gebruik feitelijke taal. Geen vage uitspraken."""

SYSTEM_PROMPT_TABLES = """Je bent een IT-adviseur. Geef ALLEEN een geldig JSON-object terug, zonder markdown, zonder uitleg.

Het JSON-object heeft twee sleutels:

"open_source_alternatieven": lijst van objecten voor closed-source producten waarvoor een reëel open-source alternatief bestaat:
  { "product": "...", "vendor": "...", "categorie": "...", "alternatief": "...", "toelichting": "..." }

"eu_alternatieven": lijst van objecten voor niet-EU producten waarvoor een EU-gebaseerd alternatief bestaat:
  { "product": "...", "vendor": "...", "regio": "...", "categorie": "...", "alternatief": "...", "toelichting": "..." }

Sla producten over waarvoor geen goed alternatief bekend is. Gebruik je eigen marktkennis."""


# ── OpenAI calls ──────────────────────────────────────────────────────────────

def _chat(system: str, user: str, max_tokens: int) -> str:
    client = OpenAI()
    response = client.chat.completions.create(
        model="gpt-4o",
        max_tokens=max_tokens,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    )
    return response.choices[0].message.content


def generate_ai_text(doelen: str, chart_data: dict) -> str:
    user = (
        f"Klantdoelen:\n{doelen}\n\n"
        f"Dashboard data:\n{json.dumps(chart_data, ensure_ascii=False, indent=2)}"
    )
    return _chat(SYSTEM_PROMPT_NARRATIVE, user, max_tokens=3000)


def generate_alternatives(doelen: str, chart_data: dict) -> dict:
    user = (
        f"Klantdoelen:\n{doelen}\n\n"
        f"Closed-source producten:\n{json.dumps(chart_data.get('closed_source_products', []), ensure_ascii=False, indent=2)}\n\n"
        f"Niet-EU producten:\n{json.dumps(chart_data.get('non_eu_products', []), ensure_ascii=False, indent=2)}"
    )
    raw = _chat(SYSTEM_PROMPT_TABLES, user, max_tokens=2000)
    # strip accidental markdown fences
    raw = re.sub(r"^```[a-z]*\n?", "", raw.strip())
    raw = re.sub(r"\n?```$", "", raw.strip())
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"open_source_alternatieven": [], "eu_alternatieven": []}


# ── chart ─────────────────────────────────────────────────────────────────────

def _build_pie_overview(chart_data: dict) -> bytes:
    tp = chart_data["type_product"]
    br = chart_data["beursregio"]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 5))
    ax1.pie(list(tp.values()), labels=list(tp.keys()), autopct="%1.1f%%", startangle=90)
    ax1.set_title("Type product", fontsize=12)
    ax2.pie(list(br.values()), labels=list(br.keys()), autopct="%1.1f%%", startangle=90)
    ax2.set_title("Beursregio", fontsize=12)
    fig.suptitle("Overzicht installed base", fontsize=14, fontweight="bold")

    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


# ── Word helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, COLOR_HEADER)

    # data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value) if value else "—"
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if r % 2 == 0:
                _set_cell_bg(cell, COLOR_ROW_ALT)

    doc.add_paragraph()


def _add_markdown_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    for segment in re.split(r"(\*\*[^*]+\*\*)", text):
        if segment.startswith("**") and segment.endswith("**"):
            para.add_run(segment[2:-2]).bold = True
        elif segment:
            para.add_run(segment)


def _add_ai_text(doc: Document, ai_text: str) -> None:
    for line in ai_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if re.fullmatch(r"\*\*[^*]+\*\*:?", stripped):
            doc.add_heading(stripped.strip("*:").strip(), level=2)
        elif stripped.startswith("- "):
            _add_markdown_paragraph(doc, stripped[2:])
        else:
            _add_markdown_paragraph(doc, stripped)


# ── main entry point ──────────────────────────────────────────────────────────

def generate_report(doelen: str, chart_data: dict) -> bytes:
    ai_text = generate_ai_text(doelen, chart_data)
    alternatives = generate_alternatives(doelen, chart_data)
    pie_overview = _build_pie_overview(chart_data)

    doc = Document()

    # Title
    doc.add_heading("Installed Base Rapport", 0)

    # Disclaimer
    disclaimer = doc.add_paragraph(
        "Dit rapport is samengesteld op basis van geautomatiseerde analyse door kunstmatige intelligentie. "
        "De gepresenteerde inzichten en aanbevelingen dienen als leidraad en dienen te worden geverifieerd "
        "door een bevoegd adviseur alvorens er beslissingen op worden gebaseerd."
    )
    disclaimer.runs[0].italic = True
    disclaimer.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph()

    # Pie charts at the top
    doc.add_picture(BytesIO(pie_overview), width=Inches(6.5))
    doc.add_paragraph()

    # Client goals
    doc.add_heading("Klantdoelen", 1)
    doc.add_paragraph(doelen)

    # Narrative analysis
    doc.add_heading("Analyse", 1)
    _add_ai_text(doc, ai_text)

    # Open source alternatives table
    os_rows = alternatives.get("open_source_alternatieven", [])
    if os_rows:
        doc.add_heading("Open Source Alternatieven", 1)
        _add_table(
            doc,
            headers=["Product", "Huidige vendor", "Categorie", "Open source alternatief", "Toelichting"],
            rows=[[r.get("product"), r.get("vendor"), r.get("categorie"), r.get("alternatief"), r.get("toelichting")] for r in os_rows],
        )

    # EU alternatives table
    eu_rows = alternatives.get("eu_alternatieven", [])
    if eu_rows:
        doc.add_heading("EU-Gebaseerde Alternatieven", 1)
        _add_table(
            doc,
            headers=["Product", "Huidige vendor", "Regio", "Categorie", "EU-alternatief", "Toelichting"],
            rows=[[r.get("product"), r.get("vendor"), r.get("regio"), r.get("categorie"), r.get("alternatief"), r.get("toelichting")] for r in eu_rows],
        )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
